from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

def set_heading_color(heading, r, g, b):
    for run in heading.runs:
        run.font.color.rgb = RGBColor(r, g, b)

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '8B5CF6')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(180, 220, 180)
    p.paragraph_format.left_indent = Inches(0.5)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1E293B')
    p_pr.append(shd)

def create_document():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    title = doc.add_heading('TaskSense', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(109, 40, 217)

    sub = doc.add_paragraph('NLP-Based Smart Action Item Extractor from Team Discussions')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    sub2 = doc.add_paragraph('Full-Stack Application Build Guide | Step-by-Step Technical Documentation')
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].font.size = Pt(11)
    sub2.runs[0].font.italic = True

    doc.add_page_break()

    # ============================================================
    # SECTION 1: OVERVIEW
    # ============================================================
    h = doc.add_heading('1. Project Overview', 1)
    set_heading_color(h, 109, 40, 217)
    add_separator(doc)

    doc.add_paragraph(
        'TaskSense is a full-stack, NLP-powered web application designed to solve a critical problem in student and team '
        'collaboration: action items get lost in long, messy chat conversations. TaskSense reads raw chat or meeting text '
        'and automatically extracts WHO needs to DO WHAT by WHEN, and at what PRIORITY — converting chaos into structured, '
        'actionable task boards.'
    )

    # ============================================================
    # SECTION 2: ARCHITECTURE OVERVIEW
    # ============================================================
    h = doc.add_heading('2. Full-Stack Architecture', 1)
    set_heading_color(h, 109, 40, 217)
    add_separator(doc)

    doc.add_paragraph('The application is divided into four primary layers:')
    layers = [
        ('Frontend (React + Vite)', 'Handles UI rendering, user input, and displaying extracted task boards.'),
        ('Backend API (Flask)', 'Processes requests, runs the NLP engine, and communicates with the database.'),
        ('NLP Engine (Python)', 'Extracts tasks, deadlines, people, and priorities using rule-based logic.'),
        ('Database (Supabase/SQLite)', 'Stores users, teams, sessions, and historical task plans.'),
    ]
    for name, desc in layers:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(name + ': ').bold = True
        p.add_run(desc)

    # ============================================================
    # SECTION 3: CURRENT STATE
    # ============================================================
    h = doc.add_heading('3. What You Already Have (Current UI)', 1)
    set_heading_color(h, 109, 40, 217)
    add_separator(doc)

    doc.add_paragraph(
        'Your current codebase (pushed to https://github.com/collab0nyn/Aichatscheduler) contains a fully built React frontend with:'
    )
    current = [
        'A premium dark-mode UI with glassmorphism design, custom CSS variables, and smooth animations.',
        'InputSection.jsx – A text area for pasting chat messages, with an "Upload .txt" button placeholder.',
        'TaskBoard.jsx – A card-based grid to display extracted tasks with Person, Task, Deadline, and Priority.',
        'App.jsx – Main layout with a sidebar navigation (Dashboard, Action Plans, Reminders, Settings).',
        'A mock NLP extraction function showing example output for the sample input.',
    ]
    for item in current:
        doc.add_paragraph(item, style='List Bullet')

    # ============================================================
    # SECTION 4: STEP-BY-STEP FEATURES TO BUILD
    # ============================================================
    h = doc.add_heading('4. Step-by-Step: Features to Build', 1)
    set_heading_color(h, 109, 40, 217)
    add_separator(doc)

    # --- FEATURE 1 ---
    h2 = doc.add_heading('Step 1: Build the Flask Backend API', 2)
    set_heading_color(h2, 236, 72, 153)
    doc.add_paragraph('Why: The frontend currently uses a mock function. You need a real server to process text with Python NLP.')
    doc.add_paragraph('How:')

    steps_1 = [
        "Install Flask: pip install flask flask-cors",
        "Create a file app.py in a new 'backend' folder inside your project.",
        "Create a POST endpoint /api/extract that accepts JSON: { \"text\": \"...chat message...\" }",
        "The endpoint calls your NLP engine and returns: { \"tasks\": [...], \"summary\": \"...\" }",
        "Run the Flask server with: python app.py",
    ]
    for s in steps_1:
        doc.add_paragraph(s, style='List Number')

    doc.add_paragraph('Sample Flask endpoint code:')
    add_code_block(doc,
        "from flask import Flask, request, jsonify\n"
        "from flask_cors import CORS\n"
        "from nlp_engine import extract_tasks, summarize\n\n"
        "app = Flask(__name__)\nCORS(app)\n\n"
        "@app.route('/api/extract', methods=['POST'])\n"
        "def extract():\n"
        "    text = request.json.get('text', '')\n"
        "    tasks = extract_tasks(text)\n"
        "    summary = summarize(text)\n"
        "    return jsonify({'tasks': tasks, 'summary': summary})\n\n"
        "if __name__ == '__main__':\n"
        "    app.run(debug=True, port=5000)"
    )

    # --- FEATURE 2 ---
    h2 = doc.add_heading('Step 2: Build the NLP Engine (Python)', 2)
    set_heading_color(h2, 236, 72, 153)
    doc.add_paragraph('Why: This is the core intelligence of TaskSense. It extracts structured data from raw text.')
    doc.add_paragraph('How (nlp_engine.py):')
    steps_2 = [
        "Deadline Extraction: Use regex to find dates. e.g., re.findall(r'by\\s+(monday|tuesday|friday|today|tomorrow)', text, re.I)",
        "Priority Detection: Check for keywords: if any(w in text.lower() for w in ['urgent','asap','immediately']): priority='High'",
        "Task Extraction: Match action verbs (prepare, complete, submit, check, do, send, review) and extract the phrase after the verb.",
        "Person Identification: Either use simple capitalized name detection or integrate spaCy: nlp = spacy.load('en_core_web_sm'); doc.ents",
        "Summarization: Rank sentences by keyword density (count action words per sentence) and return the top 2-3 sentences.",
    ]
    for s in steps_2:
        p = doc.add_paragraph(s, style='List Number')

    doc.add_paragraph('Install spaCy (optional but recommended):')
    add_code_block(doc, "pip install spacy\npython -m spacy download en_core_web_sm")

    # --- FEATURE 3 ---
    h2 = doc.add_heading('Step 3: Connect Frontend to Backend', 2)
    set_heading_color(h2, 236, 72, 153)
    doc.add_paragraph('Why: Replace the mock extraction with a real API call.')
    doc.add_paragraph('How:')
    steps_3 = [
        "In your InputSection.jsx, replace the setTimeout mock with a fetch() call.",
        "Call: const res = await fetch('http://localhost:5000/api/extract', { method:'POST', body: JSON.stringify({text: inputText}), headers:{'Content-Type':'application/json'} })",
        "Parse the response and pass the tasks and summary up to App.jsx via the onExtract callback.",
        "Handle loading and error states in the UI (show a spinner while the API is processing).",
    ]
    for s in steps_3:
        doc.add_paragraph(s, style='List Number')

    # --- FEATURE 4 ---
    h2 = doc.add_heading('Step 4: File Upload (.txt Meeting Notes)', 2)
    set_heading_color(h2, 236, 72, 153)
    doc.add_paragraph('Why: Users should be able to upload text files directly instead of only copy-pasting.')
    doc.add_paragraph('How (Frontend):')
    steps_4 = [
        "Add an <input type='file' accept='.txt'> inside InputSection.jsx.",
        "On file selection, use the FileReader API to read the file content as text.",
        "Populate the text area with the file contents, then call the extract API the same way.",
        "Add drag-and-drop support by listening to onDrop events on the textarea wrapper div.",
    ]
    for s in steps_4:
        doc.add_paragraph(s, style='List Number')

    # --- FEATURE 5 ---
    h2 = doc.add_heading('Step 5: Database Integration (Supabase)', 2)
    set_heading_color(h2, 236, 72, 153)
    doc.add_paragraph('Why: Save action plans so teams can view their history. Supabase is free and has a simple API.')
    doc.add_paragraph('How:')
    steps_5 = [
        "Create a free account at supabase.com and create a new project.",
        "Create two tables: 'action_plans' (id, title, raw_text, summary, created_at) and 'tasks' (id, plan_id, person, task, deadline, priority, is_completed).",
        "Install supabase-py: pip install supabase",
        "In your Flask backend, after extracting tasks, insert the results into Supabase.",
        "Add a GET /api/plans endpoint to retrieve historical action plans.",
        "In the frontend, add an 'Action Plans' page that fetches and lists saved plans.",
    ]
    for s in steps_5:
        doc.add_paragraph(s, style='List Number')

    # --- FEATURE 6 ---
    h2 = doc.add_heading('Step 6: User Authentication (Team Login)', 2)
    set_heading_color(h2, 236, 72, 153)
    doc.add_paragraph('Why: Allow different teams to have their own private dashboards.')
    doc.add_paragraph('How:')
    steps_6 = [
        "Use Supabase Auth (already included in your Supabase project). It handles email/password sign-up and login for free.",
        "Install the Supabase JS client in the frontend: npm install @supabase/supabase-js",
        "Create a Login and Signup page in React.",
        "Store the user session in React state (or localStorage) and protect the Dashboard route so only logged-in users can access it.",
        "Attach the user's ID to all saved plans in the database so each user/team sees only their data.",
    ]
    for s in steps_6:
        doc.add_paragraph(s, style='List Number')

    # --- FEATURE 7 ---
    h2 = doc.add_heading('Step 7: Mark Tasks as Completed', 2)
    set_heading_color(h2, 236, 72, 153)
    doc.add_paragraph('Why: Turn the task board into an actionable to-do list, not just a display.')
    doc.add_paragraph('How:')
    steps_7 = [
        "Add a checkbox or a 'Mark Complete' button to each task card in TaskBoard.jsx.",
        "On click, send a PATCH /api/tasks/:id request to the backend to update is_completed = true.",
        "In the Supabase database, update the is_completed column for that task row.",
        "Update the task card's visual style (e.g., strikethrough text, green checkmark) to reflect completion.",
    ]
    for s in steps_7:
        doc.add_paragraph(s, style='List Number')

    # --- FEATURE 8 ---
    h2 = doc.add_heading('Step 8: Export to PDF and CSV', 2)
    set_heading_color(h2, 236, 72, 153)
    doc.add_paragraph('Why: Teams need to share action plans in standard formats.')
    doc.add_paragraph('How (CSV):')
    steps_8a = [
        "In TaskBoard.jsx, on 'Export to CSV' click, convert the tasks array to a CSV string using JavaScript.",
        "Create a Blob object and trigger a download using a dynamically created <a> link.",
    ]
    for s in steps_8a:
        doc.add_paragraph(s, style='List Number')

    add_code_block(doc,
        "const csv = ['Person,Task,Deadline,Priority', ...tasks.map(t => \n"
        "  `${t.person},${t.task},${t.deadline},${t.priority}`)].join('\\n');\n"
        "const blob = new Blob([csv], {type:'text/csv'});\n"
        "const url = URL.createObjectURL(blob);\n"
        "const a = document.createElement('a'); a.href=url; a.download='action_plan.csv'; a.click();"
    )

    doc.add_paragraph('How (PDF):')
    steps_8b = [
        "Install jsPDF: npm install jspdf jspdf-autotable",
        "Import jsPDF and autoTable, then generate a styled PDF table with the task data and download it.",
    ]
    for s in steps_8b:
        doc.add_paragraph(s, style='List Number')

    # --- FEATURE 9 ---
    h2 = doc.add_heading('Step 9: Send Reminders (Email Notifications)', 2)
    set_heading_color(h2, 236, 72, 153)
    doc.add_paragraph('Why: Automatically remind team members of their tasks by email.')
    doc.add_paragraph('How:')
    steps_9 = [
        "Use the free tier of SendGrid or Resend (resend.com) for sending emails from your Flask backend.",
        "Add an email field to each task (or user profile).",
        "Create a /api/send-reminder Flask endpoint that sends an email using the provider's API.",
        "Optional: Use APScheduler (pip install apscheduler) to automatically send reminder emails the day before a deadline.",
    ]
    for s in steps_9:
        doc.add_paragraph(s, style='List Number')

    # ============================================================
    # SECTION 5: DEPLOYMENT
    # ============================================================
    h = doc.add_heading('5. Deployment', 1)
    set_heading_color(h, 109, 40, 217)
    add_separator(doc)

    deployments = [
        ('Frontend', 'Deploy for FREE on Vercel (vercel.com). Connect your GitHub repository and it deploys automatically on every push.'),
        ('Backend', 'Deploy for FREE on Render (render.com). Connect your GitHub repository, set the start command to python app.py.'),
        ('Database', 'Supabase is already cloud-hosted. No extra deployment needed.'),
    ]
    for name, desc in deployments:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(name + ': ').bold = True
        p.add_run(desc)

    # ============================================================
    # SECTION 6: RECOMMENDED BUILD ORDER
    # ============================================================
    h = doc.add_heading('6. Recommended Build Order', 1)
    set_heading_color(h, 109, 40, 217)
    add_separator(doc)

    order = [
        "Build the NLP engine in Python and test it locally with print statements.",
        "Wrap the NLP engine in a Flask API and test it with a tool like Postman.",
        "Connect the React frontend to the Flask API (replace the mock).",
        "Add file upload (.txt) support.",
        "Add database integration (Supabase) to save and retrieve plans.",
        "Add user authentication (Login/Signup).",
        "Add 'Mark as Completed' functionality.",
        "Add PDF/CSV export.",
        "Add email reminders.",
        "Deploy frontend on Vercel and backend on Render.",
    ]
    for i, step in enumerate(order, 1):
        p = doc.add_paragraph(f'Phase {i}: {step}', style='List Number')

    # ============================================================
    # SECTION 7: TOOLS AND LIBRARIES
    # ============================================================
    h = doc.add_heading('7. Tools & Libraries Reference', 1)
    set_heading_color(h, 109, 40, 217)
    add_separator(doc)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light List Accent 5'
    hdr = table.rows[0].cells
    hdr[0].text = 'Layer'
    hdr[1].text = 'Technology'
    hdr[2].text = 'Purpose'
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    rows_data = [
        ('Frontend', 'React + Vite', 'UI framework (already set up)'),
        ('Frontend', 'lucide-react', 'Icons (already installed)'),
        ('Frontend', 'jsPDF + autotable', 'PDF export'),
        ('Frontend', '@supabase/supabase-js', 'Auth + DB client'),
        ('Backend', 'Flask', 'REST API server'),
        ('Backend', 'flask-cors', 'Allow frontend to call backend'),
        ('NLP', 'spaCy', 'Named entity recognition'),
        ('NLP', 're (built-in)', 'Regex for dates/deadlines'),
        ('Database', 'Supabase', 'PostgreSQL + Auth + Storage'),
        ('Email', 'Resend / SendGrid', 'Email reminders'),
        ('Deployment', 'Vercel', 'Frontend hosting'),
        ('Deployment', 'Render', 'Backend hosting'),
    ]
    for row_data in rows_data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()
    doc.add_paragraph('Good luck building TaskSense! 🚀', style='Intense Quote')

    # Save
    doc.save('TaskSense_Full_Build_Guide.docx')
    print("Word document generated: TaskSense_Full_Build_Guide.docx")

if __name__ == '__main__':
    create_document()
